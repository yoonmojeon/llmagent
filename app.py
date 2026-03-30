"""
AI Agent — Excel 자동화 · 웹 검색 · 데이터 시각화 · 해양 도메인 RAG · 이메일 발송
LangGraph ReAct + 직접 도구 (모든 도구 in-process, MCP 서버 없음)
"""

import asyncio
import concurrent.futures
import json
import os
import re
import smtplib
import threading
import warnings
from datetime import datetime
from email import encoders
from email.header import Header
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import openpyxl
import pandas as pd

# LangGraph deprecation 경고 억제 (langgraph 1.1에서 경고만 있고 동작은 정상)
warnings.filterwarnings("ignore", message=".*create_react_agent.*")
warnings.filterwarnings("ignore", category=DeprecationWarning, module="langgraph")

import streamlit as st
from dotenv import load_dotenv
from langchain_core.messages import AIMessage, HumanMessage
from langchain_core.tools import StructuredTool
from langchain_openai import ChatOpenAI
from langgraph.prebuilt import create_react_agent
from pydantic import BaseModel, Field


load_dotenv(Path(__file__).parent / ".env")

# ─────────────────────────────────────────
# 설정
# ─────────────────────────────────────────
BASE_DIR        = Path(__file__).parent
CHROMA_DIR      = str(BASE_DIR / "rag" / "chroma_db")
COLLECTION_NAME = "maritime_knowledge"
EMBED_MODEL     = "text-embedding-3-small"
OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY")

# 모델 프로바이더 설정
MODEL_PROVIDER   = os.getenv("MODEL_PROVIDER", "openai").lower()
OPENAI_MODEL     = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
LOCAL_MODEL      = os.getenv("LOCAL_MODEL", "deepseek-r1:8b")
OLLAMA_BASE_URL  = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434/v1")
MODEL_NAME       = OPENAI_MODEL if MODEL_PROVIDER == "openai" else LOCAL_MODEL


def build_llm(provider: str | None = None) -> ChatOpenAI:
    """provider: 'openai' | 'local' | None(=.env 설정 사용)"""
    p = (provider or MODEL_PROVIDER).lower()
    if p == "local":
        return ChatOpenAI(
            model=LOCAL_MODEL,
            base_url=OLLAMA_BASE_URL,
            api_key="ollama",
            temperature=0,
        )
    return ChatOpenAI(
        model=OPENAI_MODEL,
        temperature=0,
        openai_api_key=OPENAI_API_KEY,
    )

SYSTEM_PROMPT = """당신은 생산성 향상과 해양 도메인 전문 지식을 갖춘 기업용 AI 어시스턴트입니다.
사용자 요청에 따라 Excel 자동화, 웹 검색, 데이터 시각화, 파일 관리, 해양 기술 지식 검색, 이메일 발송, SQL 데이터베이스 분석, 문서(PDF/Word) 처리 작업을 수행합니다.
항상 한국어로 응답하고, 작업 완료 후 간결하게 결과를 요약합니다.
파일 경로를 명시하지 않으면 현재 디렉토리에 적절한 이름으로 저장합니다.

사용 가능한 도구 목록:

[Excel 도구]
- create_excel, read_excel, update_cell, get_cell_value
- get_sheet_names, add_sheet, rename_sheet
- write_data_to_sheet, delete_rows, apply_formula

[유틸리티 도구]
- web_search          : DuckDuckGo 웹 검색
- get_current_datetime: 현재 날짜/시각
- list_files          : 디렉토리 파일 목록
- read_text_file      : 텍스트/CSV/JSON 파일 읽기
- create_chart        : 차트 생성 (bar/line/pie) → PNG 저장

[해양 도메인 RAG 도구]
- rag_search          : OSP/FMI 표준, DP, 크레인 선박, 해양 전력 등 지식베이스 검색
- rag_info            : 지식베이스 정보 조회

[이메일 도구]  ← Gmail SMTP, 보내는 사람 자동 설정
- send_email                  : 이메일 발송 (텍스트)
- send_email_with_attachment  : 파일 첨부 이메일 발송 (Excel, PDF, PNG 등)

[SQL 데이터베이스 도구]  ← SQLite, 기본 DB: data.db
- sql_query      : SELECT 쿼리 실행 및 결과 반환 (조회/집계/분석)
- sql_execute    : CREATE TABLE / INSERT / UPDATE / DELETE 실행
- sql_schema     : 테이블 목록 및 컬럼 정의 조회
- sql_from_excel : Excel 파일을 SQLite 테이블로 가져오기

[문서 처리 도구]  ← PDF / Word(.docx)
- read_pdf  : PDF 파일 텍스트 추출 (페이지 범위 지정 가능)
- read_word : Word(.docx) 텍스트 및 표 추출
- doc_info  : PDF/Word 메타데이터 조회 (페이지 수, 작성자 등)

[Instagram 도구]  ← DALL-E 3 AI 이미지 생성 + instagrapi 포스팅
- generate_ai_image        : DALL-E 3로 AI 이미지 생성 (글귀 합성 옵션 포함)
- add_text_to_image        : 기존 사진 위에 글귀/텍스트 합성 (PNG/JPG)
- create_text_image        : 단색 배경 감성 이미지 생성 (AI 불필요 시)
- create_carousel_images   : PPT처럼 여러 장 슬라이드 카드 이미지 일괄 생성
- instagram_post_carousel  : 여러 이미지(최대 10장)를 캐러셀 게시물로 포스팅
- instagram_post           : 단일 이미지 + 캡션 + 해시태그로 Instagram 포스팅
- instagram_check_login    : Instagram 로그인 상태 및 계정 정보 확인
- instagram_generate_hashtags : 주제별 해시태그 자동 추천

해양 기술 관련 질문(OSP, FMI, DP, 크레인 선박, 전력 관리 등)이 들어오면
반드시 rag_search를 먼저 호출하여 관련 문서를 검색한 뒤 답변하세요.

이메일 요청 시:
- 받는 사람(to), 제목(subject), 내용(body)을 반드시 확인하세요.
- 보내는 사람은 시스템에 등록된 Gmail 계정이 자동으로 사용됩니다.
- 파일 첨부가 필요하면 send_email_with_attachment를 사용하세요.

SQL 요청 시:
- 기본 DB는 data.db이며, 다른 경로를 지정할 수 있습니다.
- Excel 데이터를 SQL로 분석하고 싶으면 sql_from_excel로 먼저 가져오세요.

문서 처리 요청 시:
- PDF/Word 파일 경로를 알려주시면 텍스트를 추출하여 요약·분석합니다.

Instagram 포스팅 요청 시:
- 단일 이미지: generate_ai_image → (add_text_to_image) → instagram_post
- 캐러셀/슬라이드 여러 장: create_carousel_images → instagram_post_carousel
  · create_carousel_images에 title(커버 제목)과 slides(각 슬라이드 텍스트 리스트)를 전달하세요.
  · 반환된 파일 경로 목록을 instagram_post_carousel의 image_paths에 그대로 넣으세요.
- instagram_generate_hashtags로 해시태그를 추천받아 캡션에 포함하세요.
- 포스팅 전 instagram_check_login으로 로그인 상태를 확인하는 것을 권장합니다."""


# ─────────────────────────────────────────
# Excel 도구 (직접 구현)
# ─────────────────────────────────────────

def _excel_read(file_path: str, sheet_name: str = "", max_rows: int = 100) -> str:
    """Excel 파일을 읽어 내용을 반환합니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    sheet = sheet_name if sheet_name else 0
    df = pd.read_excel(file_path, sheet_name=sheet, nrows=max_rows)
    return f"파일: {file_path}\n크기: {df.shape[0]}행 × {df.shape[1]}열\n\n{df.to_string(index=True)}"


def _excel_create(file_path: str, headers: list[str] = [], data: list[list] = [], sheet_name: str = "Sheet1") -> str:
    """새 Excel 파일을 생성하고 데이터를 입력합니다."""
    os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    if headers:
        ws.append(headers)
    for row in data:
        ws.append(row)
    wb.save(file_path)
    row_count = len(data)
    col_count = len(headers) if headers else (len(data[0]) if data else 0)
    return f"Excel 파일 생성 완료: {file_path}\n시트: {sheet_name}\n데이터: {row_count}행 × {col_count}열"


def _excel_update_cell(file_path: str, cell: str, value: str, sheet_name: str = "") -> str:
    """Excel 파일의 특정 셀 값을 업데이트합니다."""
    if not os.path.exists(file_path):
        wb = openpyxl.Workbook(); ws = wb.active
    else:
        wb = openpyxl.load_workbook(file_path)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    ws[cell] = value
    wb.save(file_path)
    return f"셀 {cell} → '{value}' 업데이트 완료 (파일: {file_path})"


def _excel_get_cell(file_path: str, cell: str, sheet_name: str = "") -> str:
    """Excel 파일에서 특정 셀의 값을 읽어옵니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    wb = openpyxl.load_workbook(file_path)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    return f"셀 {cell} 값: {ws[cell].value}"


def _excel_get_sheets(file_path: str) -> str:
    """Excel 파일에 있는 모든 시트 이름을 반환합니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    wb = openpyxl.load_workbook(file_path)
    sheets = wb.sheetnames
    return f"시트 목록 ({len(sheets)}개): {', '.join(sheets)}"


def _excel_add_sheet(file_path: str, sheet_name: str) -> str:
    """Excel 파일에 새 시트를 추가합니다."""
    if not os.path.exists(file_path):
        wb = openpyxl.Workbook(); wb.active.title = sheet_name
    else:
        wb = openpyxl.load_workbook(file_path)
        if sheet_name in wb.sheetnames:
            return f"시트 '{sheet_name}'이 이미 존재합니다."
        wb.create_sheet(sheet_name)
    wb.save(file_path)
    return f"시트 '{sheet_name}' 추가 완료 (파일: {file_path})"


def _excel_rename_sheet(file_path: str, old_name: str, new_name: str) -> str:
    """Excel 파일의 시트 이름을 변경합니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    wb = openpyxl.load_workbook(file_path)
    if old_name not in wb.sheetnames:
        return f"시트 '{old_name}'을 찾을 수 없습니다."
    wb[old_name].title = new_name
    wb.save(file_path)
    return f"시트 이름 변경: '{old_name}' → '{new_name}'"


def _excel_write_data(file_path: str, sheet_name: str, data: list[list],
                      start_row: int = 1, start_col: int = 1) -> str:
    """Excel 파일의 특정 시트에 데이터를 일괄 입력합니다."""
    if not os.path.exists(file_path):
        wb = openpyxl.Workbook(); wb.active.title = sheet_name
    else:
        wb = openpyxl.load_workbook(file_path)
    if sheet_name not in wb.sheetnames:
        wb.create_sheet(sheet_name)
    ws = wb[sheet_name]
    for r_idx, row in enumerate(data):
        for c_idx, val in enumerate(row):
            ws.cell(row=start_row + r_idx, column=start_col + c_idx, value=val)
    wb.save(file_path)
    return f"{len(data)}행 데이터를 '{sheet_name}' 시트에 입력 완료"


def _excel_delete_rows(file_path: str, row_start: int, row_end: int = 0, sheet_name: str = "") -> str:
    """Excel 파일에서 특정 행(들)을 삭제합니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    end = row_end if row_end else row_start
    wb = openpyxl.load_workbook(file_path)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    ws.delete_rows(row_start, end - row_start + 1)
    wb.save(file_path)
    return f"행 {row_start}~{end} 삭제 완료"


def _excel_apply_formula(file_path: str, cell: str, formula: str, sheet_name: str = "") -> str:
    """Excel 셀에 수식을 적용합니다."""
    if not os.path.exists(file_path):
        wb = openpyxl.Workbook(); ws = wb.active
    else:
        wb = openpyxl.load_workbook(file_path)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    ws[cell] = formula
    wb.save(file_path)
    return f"수식 '{formula}' → 셀 {cell} 적용 완료"


class _ExcelReadInput(BaseModel):
    file_path: str = Field(description="읽을 Excel 파일 경로")
    sheet_name: str = Field(default="", description="시트 이름 (생략 시 첫 번째 시트)")
    max_rows: int = Field(default=100, description="최대 읽을 행 수")

class _ExcelCreateInput(BaseModel):
    file_path: str = Field(description="생성할 Excel 파일 경로")
    headers: list[str] = Field(default=[], description="헤더 행 (컬럼명 목록)")
    data: list[list] = Field(default=[], description="데이터 2D 배열")
    sheet_name: str = Field(default="Sheet1", description="시트 이름")

class _ExcelCellInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    cell: str = Field(description="셀 주소 (예: A1, B3)")
    sheet_name: str = Field(default="", description="시트 이름")

class _ExcelUpdateInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    cell: str = Field(description="셀 주소 (예: A1)")
    value: str = Field(description="입력할 값")
    sheet_name: str = Field(default="", description="시트 이름")

class _ExcelSheetInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    sheet_name: str = Field(description="시트 이름")

class _ExcelRenameInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    old_name: str = Field(description="현재 시트 이름")
    new_name: str = Field(description="새 시트 이름")

class _ExcelWriteInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    sheet_name: str = Field(description="시트 이름")
    data: list[list] = Field(description="입력할 데이터 2D 배열")
    start_row: int = Field(default=1, description="시작 행 번호")
    start_col: int = Field(default=1, description="시작 열 번호")

class _ExcelDeleteRowsInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    row_start: int = Field(description="삭제 시작 행 번호")
    row_end: int = Field(default=0, description="삭제 종료 행 번호 (0이면 row_start만)")
    sheet_name: str = Field(default="", description="시트 이름")

class _ExcelFormulaInput(BaseModel):
    file_path: str = Field(description="Excel 파일 경로")
    cell: str = Field(description="수식 입력할 셀 주소")
    formula: str = Field(description="Excel 수식 (예: =SUM(A1:A9))")
    sheet_name: str = Field(default="", description="시트 이름")

EXCEL_TOOLS = [
    StructuredTool.from_function(func=_excel_read,         name="read_excel",        description="Excel 파일을 읽어 내용을 반환합니다.",                              args_schema=_ExcelReadInput),
    StructuredTool.from_function(func=_excel_create,       name="create_excel",      description="새 Excel 파일을 생성하고 데이터를 입력합니다.",                     args_schema=_ExcelCreateInput),
    StructuredTool.from_function(func=_excel_update_cell,  name="update_cell",       description="Excel 파일의 특정 셀 값을 업데이트합니다.",                         args_schema=_ExcelUpdateInput),
    StructuredTool.from_function(func=_excel_get_cell,     name="get_cell_value",    description="Excel 파일에서 특정 셀의 값을 읽어옵니다.",                         args_schema=_ExcelCellInput),
    StructuredTool.from_function(func=_excel_get_sheets,   name="get_sheet_names",   description="Excel 파일에 있는 모든 시트 이름을 반환합니다.",                    args_schema=_ExcelSheetInput),
    StructuredTool.from_function(func=_excel_add_sheet,    name="add_sheet",         description="Excel 파일에 새 시트를 추가합니다.",                                args_schema=_ExcelSheetInput),
    StructuredTool.from_function(func=_excel_rename_sheet, name="rename_sheet",      description="Excel 파일의 시트 이름을 변경합니다.",                              args_schema=_ExcelRenameInput),
    StructuredTool.from_function(func=_excel_write_data,   name="write_data_to_sheet", description="Excel 파일의 특정 시트에 데이터를 일괄 입력합니다.",              args_schema=_ExcelWriteInput),
    StructuredTool.from_function(func=_excel_delete_rows,  name="delete_rows",       description="Excel 파일에서 특정 행(들)을 삭제합니다.",                          args_schema=_ExcelDeleteRowsInput),
    StructuredTool.from_function(func=_excel_apply_formula,name="apply_formula",     description="Excel 셀에 수식(=SUM, =AVERAGE 등)을 적용합니다.",                  args_schema=_ExcelFormulaInput),
]


# ─────────────────────────────────────────
# Utils 도구 (직접 구현)
# ─────────────────────────────────────────

def _web_search(query: str, max_results: int = 5) -> str:
    """DuckDuckGo로 웹을 검색합니다."""
    from duckduckgo_search import DDGS
    results = list(DDGS().text(query, max_results=max_results))
    if not results:
        return f"'{query}' 에 대한 검색 결과가 없습니다."
    lines = [f"검색어: {query}  |  결과 {len(results)}건\n"]
    for i, r in enumerate(results, 1):
        lines.append(f"[{i}] {r.get('title', '')}")
        lines.append(f"    {r.get('href', '')}")
        body = r.get("body", "")
        if body:
            lines.append(f"    {body[:250]}")
        lines.append("")
    return "\n".join(lines)


def _get_datetime() -> str:
    """현재 날짜와 시각을 반환합니다."""
    now = datetime.now()
    weekdays = {"Monday":"월요일","Tuesday":"화요일","Wednesday":"수요일",
                "Thursday":"목요일","Friday":"금요일","Saturday":"토요일","Sunday":"일요일"}
    day_kor = weekdays.get(now.strftime("%A"), now.strftime("%A"))
    return f"현재 날짜/시각: {now.strftime('%Y년 %m월 %d일')} ({day_kor}) {now.strftime('%H:%M:%S')}"


def _list_files(directory_path: str, extension: str = "") -> str:
    """지정한 디렉토리 안의 파일과 폴더 목록을 반환합니다."""
    if not os.path.exists(directory_path):
        return f"디렉토리를 찾을 수 없습니다: {directory_path}"
    entries = []
    for entry in sorted(Path(directory_path).iterdir()):
        if extension and not entry.name.endswith(extension):
            continue
        if entry.is_file():
            size = entry.stat().st_size
            entries.append(f"  [파일] {entry.name:<40} {size/1024:.1f} KB" if size >= 1024 else f"  [파일] {entry.name:<40} {size} B")
        else:
            entries.append(f"  [폴더] {entry.name}/")
    if not entries:
        return f"'{directory_path}' 에 항목이 없습니다."
    return f"디렉토리: {directory_path}  ({len(entries)}개)\n" + "\n".join(entries)


def _read_text_file(file_path: str, max_lines: int = 100) -> str:
    """텍스트, CSV, JSON 파일을 읽어 내용을 반환합니다."""
    if not os.path.exists(file_path):
        return f"파일을 찾을 수 없습니다: {file_path}"
    ext = Path(file_path).suffix.lower()
    if ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return f"JSON: {file_path}\n\n{json.dumps(data, ensure_ascii=False, indent=2)}"
    if ext == ".csv":
        df = pd.read_csv(file_path, nrows=max_lines)
        return f"CSV: {file_path}  ({df.shape[0]}행 × {df.shape[1]}열)\n\n{df.to_string()}"
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        lines = f.readlines()
    total = len(lines)
    content = "".join(lines[:max_lines])
    suffix = f"\n\n... ({total}줄 중 {max_lines}줄 표시)" if total > max_lines else ""
    return f"파일: {file_path}\n\n{content}{suffix}"


def _create_chart(chart_type: str, title: str, labels: list[str], values: list[float],
                  save_path: str, x_label: str = "", y_label: str = "") -> str:
    """막대(bar), 선(line), 원형(pie) 차트를 생성하고 PNG로 저장합니다."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams["font.family"] = "Malgun Gothic"
    plt.rcParams["axes.unicode_minus"] = False
    PALETTE = ["#2D3748","#553C9A","#2B6CB0","#276749","#C05621","#702459","#1A365D","#44337A","#1C4532","#7B341E"]
    fig, ax = plt.subplots(figsize=(11, 6))
    fig.patch.set_facecolor("white"); ax.set_facecolor("#FAFAFA")
    if chart_type == "bar":
        bars = ax.bar(labels, values, color=PALETTE[:len(labels)], edgecolor="white", linewidth=0.8, width=0.6)
        for bar, val in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(values)*0.012,
                    f"{val:,.0f}", ha="center", va="bottom", fontsize=9, color="#2D3748")
        ax.grid(axis="y", alpha=0.25, linestyle="--", color="#A0AEC0"); ax.set_axisbelow(True)
        if x_label: ax.set_xlabel(x_label, fontsize=10)
        if y_label: ax.set_ylabel(y_label, fontsize=10)
    elif chart_type == "line":
        ax.plot(range(len(labels)), values, color=PALETTE[0], marker="o", linewidth=2.5,
                markersize=7, markerfacecolor="white", markeredgewidth=2)
        ax.fill_between(range(len(labels)), values, alpha=0.08, color=PALETTE[0])
        ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels)
        ax.grid(alpha=0.25, linestyle="--", color="#A0AEC0"); ax.set_axisbelow(True)
        if x_label: ax.set_xlabel(x_label, fontsize=10)
        if y_label: ax.set_ylabel(y_label, fontsize=10)
    elif chart_type == "pie":
        wedges, texts, autotexts = ax.pie(values, labels=labels, autopct="%1.1f%%",
            colors=PALETTE[:len(labels)], startangle=90,
            wedgeprops={"edgecolor":"white","linewidth":2}, pctdistance=0.82)
        for at in autotexts:
            at.set_fontsize(9); at.set_color("white")
    for spine in ["top","right"]: ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#E2E8F0"); ax.spines["bottom"].set_color("#E2E8F0")
    ax.tick_params(colors="#4A5568")
    ax.set_title(title, fontsize=15, fontweight="bold", color="#1A202C", pad=16)
    plt.tight_layout()
    os.makedirs(os.path.dirname(os.path.abspath(save_path)), exist_ok=True)
    plt.savefig(save_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()
    return f"차트 저장 완료: {save_path}"


class _WebSearchInput(BaseModel):
    query: str = Field(description="검색할 키워드 또는 문장")
    max_results: int = Field(default=5, description="반환할 최대 결과 수")

class _ListFilesInput(BaseModel):
    directory_path: str = Field(description="조회할 디렉토리 경로")
    extension: str = Field(default="", description="필터링할 확장자 (예: .xlsx)")

class _ReadTextInput(BaseModel):
    file_path: str = Field(description="읽을 파일 경로 (.txt/.csv/.json)")
    max_lines: int = Field(default=100, description="최대 읽을 줄 수")

class _ChartInput(BaseModel):
    chart_type: str = Field(description="차트 종류: bar(막대), line(선), pie(원형)")
    title: str = Field(description="차트 제목")
    labels: list[str] = Field(description="레이블 목록")
    values: list[float] = Field(description="수치 데이터 목록")
    save_path: str = Field(description="저장할 PNG 파일 경로")
    x_label: str = Field(default="", description="X축 이름")
    y_label: str = Field(default="", description="Y축 이름")

UTILS_TOOLS = [
    StructuredTool.from_function(func=_web_search,      name="web_search",          description="DuckDuckGo로 인터넷 검색합니다. 최신 뉴스, 일반 지식 등을 찾을 때 사용.", args_schema=_WebSearchInput),
    StructuredTool.from_function(func=_get_datetime,    name="get_current_datetime", description="현재 날짜와 시각을 반환합니다."),
    StructuredTool.from_function(func=_list_files,      name="list_files",           description="디렉토리 안의 파일과 폴더 목록을 반환합니다.",                            args_schema=_ListFilesInput),
    StructuredTool.from_function(func=_read_text_file,  name="read_text_file",       description="텍스트·CSV·JSON 파일을 읽어 내용을 반환합니다.",                           args_schema=_ReadTextInput),
    StructuredTool.from_function(func=_create_chart,    name="create_chart",         description="막대/선/원형 차트를 생성하고 PNG로 저장합니다.",                            args_schema=_ChartInput),
]


# ─────────────────────────────────────────
# Email 도구 (직접 구현)
# ─────────────────────────────────────────

def _smtp_connect() -> smtplib.SMTP:
    server = smtplib.SMTP("smtp.gmail.com", 587)
    server.ehlo(); server.starttls()
    server.login(os.getenv("GMAIL_EMAIL", ""), os.getenv("GMAIL_APP_PASSWORD", ""))
    return server


def _email_check_config() -> str | None:
    if not os.getenv("GMAIL_EMAIL"):
        return "오류: .env에 GMAIL_EMAIL이 설정되지 않았습니다."
    if not os.getenv("GMAIL_APP_PASSWORD"):
        return "Gmail 앱 비밀번호가 설정되지 않았습니다. myaccount.google.com > 보안 > 앱 비밀번호"
    return None


def _build_mime(to: str, subject: str, body: str, cc: str = "") -> tuple[MIMEMultipart, list[str], str]:
    from_addr = os.getenv("GMAIL_EMAIL", "")
    msg = MIMEMultipart()
    msg["From"] = from_addr; msg["To"] = to
    msg["Subject"] = str(Header(subject, "utf-8"))
    if cc: msg["Cc"] = cc
    msg.attach(MIMEText(body, "plain", "utf-8"))
    recipients = [t.strip() for t in to.split(",")]
    if cc: recipients += [c.strip() for c in cc.split(",")]
    return msg, recipients, from_addr


def _send_email(to: str, subject: str, body: str, cc: str = "") -> str:
    """Gmail SMTP로 이메일을 발송합니다."""
    err = _email_check_config()
    if err: return err
    msg, recipients, from_addr = _build_mime(to, subject, body, cc)
    try:
        with _smtp_connect() as server:
            server.sendmail(from_addr, recipients, msg.as_string())
        return f"이메일 발송 완료\n보낸 사람: {from_addr}\n받는 사람: {to}\n제목: {subject}"
    except smtplib.SMTPAuthenticationError:
        return "인증 실패: 앱 비밀번호를 확인하세요. (myaccount.google.com > 보안 > 앱 비밀번호)"
    except Exception as e:
        return f"발송 실패: {e}"


def _send_email_attachment(to: str, subject: str, body: str, file_path: str, cc: str = "") -> str:
    """파일을 첨부하여 Gmail SMTP로 이메일을 발송합니다."""
    err = _email_check_config()
    if err: return err
    if not Path(file_path).exists():
        return f"오류: 파일을 찾을 수 없습니다 — {file_path}"
    msg, recipients, from_addr = _build_mime(to, subject, body, cc)
    filename = Path(file_path).name
    with open(file_path, "rb") as f:
        part = MIMEBase("application", "octet-stream"); part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", "attachment", filename=("utf-8", "", filename))
    msg.attach(part)
    file_kb = Path(file_path).stat().st_size / 1024
    try:
        with _smtp_connect() as server:
            server.sendmail(from_addr, recipients, msg.as_string())
        return f"이메일 발송 완료\n보낸 사람: {from_addr}\n받는 사람: {to}\n제목: {subject}\n첨부: {filename} ({file_kb:.1f} KB)"
    except smtplib.SMTPAuthenticationError:
        return "인증 실패: 앱 비밀번호를 확인하세요."
    except Exception as e:
        return f"발송 실패: {e}"


class _SendEmailInput(BaseModel):
    to: str = Field(description="받는 사람 이메일 주소 (여러 명은 쉼표로 구분)")
    subject: str = Field(description="이메일 제목")
    body: str = Field(description="이메일 본문")
    cc: str = Field(default="", description="참조(CC) 이메일 주소")

class _SendEmailAttachInput(BaseModel):
    to: str = Field(description="받는 사람 이메일 주소")
    subject: str = Field(description="이메일 제목")
    body: str = Field(description="이메일 본문")
    file_path: str = Field(description="첨부할 파일 경로")
    cc: str = Field(default="", description="참조(CC) 이메일 주소")

EMAIL_TOOLS = [
    StructuredTool.from_function(func=_send_email,           name="send_email",                description="Gmail SMTP로 이메일을 발송합니다. 보내는 사람은 .env의 Gmail 계정이 자동 사용됩니다.", args_schema=_SendEmailInput),
    StructuredTool.from_function(func=_send_email_attachment, name="send_email_with_attachment", description="파일을 첨부하여 Gmail SMTP로 이메일을 발송합니다.",                                   args_schema=_SendEmailAttachInput),
]


# ─────────────────────────────────────────
# RAG 싱글톤 — ChromaDB를 최초 1회만 로드
# ─────────────────────────────────────────
_rag_collection = None
_rag_lock       = threading.Lock()


def _get_rag_collection():
    """ChromaDB 컬렉션 싱글톤 (스레드 안전)"""
    global _rag_collection
    if _rag_collection is not None:
        return _rag_collection
    with _rag_lock:
        if _rag_collection is not None:
            return _rag_collection
        import chromadb
        from chromadb.utils.embedding_functions import OpenAIEmbeddingFunction
        embed_fn = OpenAIEmbeddingFunction(
            api_key=OPENAI_API_KEY,
            model_name=EMBED_MODEL,
        )
        client = chromadb.PersistentClient(path=CHROMA_DIR)
        _rag_collection = client.get_collection(
            COLLECTION_NAME, embedding_function=embed_fn
        )
        return _rag_collection


# ── RAG 도구 함수 ─────────────────────────

class _RagSearchInput(BaseModel):
    query: str = Field(description="검색할 질문 또는 키워드 (한국어/영어 모두 가능)")
    top_k: int = Field(default=4, description="반환할 최대 문서 청크 수 (기본: 4)")


class _RagInfoInput(BaseModel):
    pass


def _rag_search(query: str, top_k: int = 4) -> str:
    col = _get_rag_collection()
    n   = min(int(top_k), col.count())
    if n == 0:
        return "지식베이스가 비어 있습니다."
    results = col.query(query_texts=[query], n_results=n)
    parts = []
    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0],
    ):
        sim    = max(0.0, 1.0 - dist)
        source = meta.get("source", "unknown")
        page   = meta.get("page", "")
        loc    = source + (f" p.{page}" if page else "")
        parts.append(f"[{loc}] (유사도: {sim:.3f})\n{doc.strip()}")
    header = f'검색어: "{query}"  |  결과 {len(parts)}건\n\n'
    return header + "\n\n---\n\n".join(parts)


def _rag_info() -> str:
    col = _get_rag_collection()
    docs_dir = BASE_DIR / "rag" / "docs"
    sources  = []
    if docs_dir.exists():
        for f in sorted(docs_dir.iterdir()):
            size_str = f"{f.stat().st_size / 1024:.0f} KB"
            sources.append(f"  - {f.name}  ({size_str})")
    lines = [
        f"컬렉션: {COLLECTION_NAME}",
        f"임베딩 모델: {EMBED_MODEL}",
        f"총 청크 수: {col.count():,}개",
        "",
        "문서 소스:",
    ] + sources
    return "\n".join(lines)


RAG_TOOLS = [
    StructuredTool.from_function(
        func=_rag_search,
        name="rag_search",
        description=(
            "해양 도메인 지식베이스(OSP Interface Specification, FMI 2.0 표준, "
            "동적 위치 유지 DP, 크레인 선박, 해양 전력 관리, 해양 건설)에서 "
            "의미 유사도 기반으로 관련 문서를 검색합니다."
        ),
        args_schema=_RagSearchInput,
    ),
    StructuredTool.from_function(
        func=_rag_info,
        name="rag_info",
        description="현재 해양 지식베이스의 정보를 반환합니다 (문서 수, 소스 목록).",
        args_schema=_RagInfoInput,
    ),
]


# 이메일 도구 (직접 구현)


# ─────────────────────────────────────────
# SQL 데이터베이스 도구 — sqlite3 내장 (빠름)
# ─────────────────────────────────────────
DEFAULT_DB = str(BASE_DIR / "data.db")


class _SqlQueryInput(BaseModel):
    query:   str = Field(description="실행할 SELECT SQL 문")
    db_path: str = Field(default=DEFAULT_DB, description="SQLite DB 파일 경로 (기본: data.db)")


class _SqlExecuteInput(BaseModel):
    statement: str = Field(description="실행할 SQL 문 (CREATE TABLE, INSERT, UPDATE, DELETE 등)")
    db_path:   str = Field(default=DEFAULT_DB, description="SQLite DB 파일 경로 (기본: data.db)")


class _SqlSchemaInput(BaseModel):
    db_path: str = Field(default=DEFAULT_DB, description="SQLite DB 파일 경로 (기본: data.db)")


class _SqlFromExcelInput(BaseModel):
    excel_path: str = Field(description="불러올 Excel 파일 경로")
    table_name: str = Field(description="생성할 테이블 이름")
    db_path:    str = Field(default=DEFAULT_DB, description="SQLite DB 파일 경로 (기본: data.db)")
    sheet:      str = Field(default="", description="읽을 시트 이름 (비워두면 첫 번째 시트)")


def _sql_query(query: str, db_path: str = DEFAULT_DB) -> str:
    """SELECT 쿼리를 실행하고 결과를 테이블 형식으로 반환합니다."""
    import sqlite3
    try:
        con = sqlite3.connect(db_path)
        cur = con.execute(query)
        rows = cur.fetchall()
        cols = [d[0] for d in cur.description] if cur.description else []
        con.close()
        if not rows:
            return "쿼리 결과가 없습니다."
        col_w = [max(len(str(c)), max((len(str(r[i])) for r in rows), default=0)) for i, c in enumerate(cols)]
        header = " | ".join(str(c).ljust(col_w[i]) for i, c in enumerate(cols))
        sep    = "-+-".join("-" * w for w in col_w)
        lines  = [header, sep] + [
            " | ".join(str(r[i]).ljust(col_w[i]) for i in range(len(cols))) for r in rows
        ]
        return f"결과 {len(rows)}행\n\n" + "\n".join(lines)
    except Exception as e:
        return f"SQL 오류: {e}"


def _sql_execute(statement: str, db_path: str = DEFAULT_DB) -> str:
    """CREATE TABLE, INSERT, UPDATE, DELETE 등 DML/DDL을 실행합니다."""
    import sqlite3
    try:
        con = sqlite3.connect(db_path)
        cur = con.execute(statement)
        con.commit()
        affected = cur.rowcount
        con.close()
        msg = f"실행 완료 (영향받은 행: {affected})" if affected >= 0 else "실행 완료"
        return f"{msg}\nDB: {db_path}"
    except Exception as e:
        return f"SQL 오류: {e}"


def _sql_schema(db_path: str = DEFAULT_DB) -> str:
    """DB의 모든 테이블 목록과 컬럼 정의를 반환합니다."""
    import sqlite3
    if not Path(db_path).exists():
        return f"DB 파일이 없습니다: {db_path}"
    try:
        con = sqlite3.connect(db_path)
        tables = [r[0] for r in con.execute(
            "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name"
        ).fetchall()]
        if not tables:
            return "테이블이 없습니다."
        parts = []
        for tbl in tables:
            cols = con.execute(f"PRAGMA table_info({tbl})").fetchall()
            col_defs = ", ".join(f"{c[1]} {c[2]}" for c in cols)
            cnt = con.execute(f"SELECT COUNT(*) FROM {tbl}").fetchone()[0]
            parts.append(f"[{tbl}]  ({cnt:,}행)\n  {col_defs}")
        con.close()
        return f"DB: {db_path}\n테이블 {len(tables)}개\n\n" + "\n\n".join(parts)
    except Exception as e:
        return f"스키마 조회 오류: {e}"


def _sql_from_excel(excel_path: str, table_name: str,
                    db_path: str = DEFAULT_DB, sheet: str = "") -> str:
    """Excel 파일을 읽어 SQLite 테이블로 가져옵니다 (기존 테이블 교체)."""
    import sqlite3
    import pandas as pd
    if not Path(excel_path).exists():
        return f"오류: 파일을 찾을 수 없습니다 — {excel_path}"
    try:
        df = pd.read_excel(excel_path, sheet_name=(sheet if sheet else 0))
        con = sqlite3.connect(db_path)
        df.to_sql(table_name, con, if_exists="replace", index=False)
        con.close()
        return (
            f"Excel → SQLite 가져오기 완료\n"
            f"파일: {excel_path}\n"
            f"테이블: {table_name}  ({len(df):,}행 × {len(df.columns)}열)\n"
            f"DB: {db_path}"
        )
    except Exception as e:
        return f"가져오기 오류: {e}"


SQL_TOOLS = [
    StructuredTool.from_function(
        func=_sql_query,
        name="sql_query",
        description=(
            "SQLite DB에서 SELECT 쿼리를 실행하고 결과를 반환합니다. "
            "데이터 조회, 집계, 분석에 사용합니다."
        ),
        args_schema=_SqlQueryInput,
    ),
    StructuredTool.from_function(
        func=_sql_execute,
        name="sql_execute",
        description=(
            "SQLite DB에서 CREATE TABLE, INSERT, UPDATE, DELETE 등을 실행합니다. "
            "테이블 생성 및 데이터 변경에 사용합니다."
        ),
        args_schema=_SqlExecuteInput,
    ),
    StructuredTool.from_function(
        func=_sql_schema,
        name="sql_schema",
        description="SQLite DB의 테이블 목록, 컬럼 정의, 행 수를 반환합니다.",
        args_schema=_SqlSchemaInput,
    ),
    StructuredTool.from_function(
        func=_sql_from_excel,
        name="sql_from_excel",
        description=(
            "Excel 파일을 SQLite DB 테이블로 가져옵니다. "
            "Excel 데이터를 DB로 전환하거나 SQL 분석을 준비할 때 사용합니다."
        ),
        args_schema=_SqlFromExcelInput,
    ),
]


# ─────────────────────────────────────────
# 문서 처리 도구 — PDF / Word (.docx)
# ─────────────────────────────────────────

class _ReadPdfInput(BaseModel):
    file_path:  str = Field(description="PDF 파일 경로")
    page_start: int = Field(default=1, description="시작 페이지 번호 (1부터, 기본: 1)")
    page_end:   int = Field(default=0, description="끝 페이지 번호 (0이면 마지막 페이지까지)")


class _ReadWordInput(BaseModel):
    file_path: str = Field(description="Word(.docx) 파일 경로")


class _DocInfoInput(BaseModel):
    file_path: str = Field(description="PDF 또는 Word(.docx) 파일 경로")


def _read_pdf(file_path: str, page_start: int = 1, page_end: int = 0) -> str:
    """PDF 파일에서 텍스트를 추출합니다."""
    from pypdf import PdfReader
    if not Path(file_path).exists():
        return f"오류: 파일을 찾을 수 없습니다 — {file_path}"
    try:
        reader = PdfReader(file_path)
        total  = len(reader.pages)
        start  = max(0, page_start - 1)
        end    = total if (page_end == 0 or page_end > total) else page_end
        parts  = []
        for i in range(start, end):
            text = reader.pages[i].extract_text() or ""
            parts.append(f"--- 페이지 {i + 1} ---\n{text.strip()}")
        header = f"PDF: {Path(file_path).name}  (전체 {total}p, 추출: {start+1}~{end}p)\n\n"
        return header + "\n\n".join(parts) if parts else header + "(텍스트 없음)"
    except Exception as e:
        return f"PDF 읽기 오류: {e}"


def _read_word(file_path: str) -> str:
    """Word(.docx) 파일에서 텍스트를 추출합니다."""
    from docx import Document
    if not Path(file_path).exists():
        return f"오류: 파일을 찾을 수 없습니다 — {file_path}"
    try:
        doc   = Document(file_path)
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # 표 내 텍스트도 포함
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        header = f"Word: {Path(file_path).name}  (단락 {len(doc.paragraphs)}개, 표 {len(doc.tables)}개)\n\n"
        return header + "\n".join(paras) if paras else header + "(텍스트 없음)"
    except Exception as e:
        return f"Word 읽기 오류: {e}"


def _doc_info(file_path: str) -> str:
    """PDF 또는 Word 파일의 메타데이터를 반환합니다."""
    path = Path(file_path)
    if not path.exists():
        return f"오류: 파일을 찾을 수 없습니다 — {file_path}"
    ext = path.suffix.lower()
    size_kb = path.stat().st_size / 1024
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            meta   = reader.metadata or {}
            info   = [
                f"파일: {path.name}  ({size_kb:.1f} KB)",
                f"형식: PDF",
                f"페이지 수: {len(reader.pages)}",
                f"제목: {meta.get('/Title', '없음')}",
                f"작성자: {meta.get('/Author', '없음')}",
                f"생성일: {meta.get('/CreationDate', '없음')}",
                f"암호화: {'예' if reader.is_encrypted else '아니요'}",
            ]
        elif ext in (".docx", ".doc"):
            from docx import Document
            doc  = Document(file_path)
            prop = doc.core_properties
            info = [
                f"파일: {path.name}  ({size_kb:.1f} KB)",
                f"형식: Word",
                f"단락 수: {len(doc.paragraphs)}",
                f"표 수: {len(doc.tables)}",
                f"제목: {prop.title or '없음'}",
                f"작성자: {prop.author or '없음'}",
                f"최종 수정: {prop.modified or '없음'}",
            ]
        else:
            return f"지원하지 않는 형식입니다 (PDF 또는 .docx만 가능): {ext}"
        return "\n".join(info)
    except Exception as e:
        return f"파일 정보 조회 오류: {e}"


DOCUMENT_TOOLS = [
    StructuredTool.from_function(
        func=_read_pdf,
        name="read_pdf",
        description=(
            "PDF 파일에서 텍스트를 추출합니다. "
            "계약서, 보고서, 기술 문서 분석에 사용합니다. "
            "페이지 범위를 지정할 수 있습니다."
        ),
        args_schema=_ReadPdfInput,
    ),
    StructuredTool.from_function(
        func=_read_word,
        name="read_word",
        description=(
            "Word(.docx) 파일에서 텍스트와 표 내용을 추출합니다. "
            "문서 요약, 내용 분석, 데이터 추출에 사용합니다."
        ),
        args_schema=_ReadWordInput,
    ),
    StructuredTool.from_function(
        func=_doc_info,
        name="doc_info",
        description=(
            "PDF 또는 Word 파일의 메타데이터(페이지 수, 작성자, 생성일 등)를 반환합니다."
        ),
        args_schema=_DocInfoInput,
    ),
]


# ─────────────────────────────────────────
# Instagram 도구 — instagrapi 직접 호출
# ─────────────────────────────────────────

class _AddTextToImageInput(BaseModel):
    image_path: str = Field(description="원본 사진 파일 경로 (JPG, PNG)")
    text: str = Field(description="사진에 넣을 글귀 (줄바꿈은 \\n 사용)")
    output_path: str = Field(default="", description="저장할 파일 경로 (비워두면 원본명_text.jpg)")
    font_size: int = Field(default=48, description="글자 크기 (기본: 48)")
    text_color: str = Field(default="white", description="글자 색 — white, black, yellow, #RRGGBB")
    position: str = Field(default="center", description="텍스트 위치 — center, top, bottom")
    shadow: bool = Field(default=True, description="그림자 효과 (기본: True)")
    overlay_opacity: int = Field(default=40, description="어두운 오버레이 투명도 0~100 (기본: 40)")


class _CreateTextImageInput(BaseModel):
    text: str = Field(description="메인 글귀 (줄바꿈은 \\n 사용)")
    output_path: str = Field(default="instagram_post.jpg", description="저장할 파일 경로")
    width: int = Field(default=1080, description="이미지 너비 px")
    height: int = Field(default=1080, description="이미지 높이 px")
    bg_color: str = Field(default="#1A1A2E", description="배경색 — #RRGGBB 또는 black, white, navy, purple")
    text_color: str = Field(default="white", description="글자 색상")
    font_size: int = Field(default=60, description="메인 글자 크기")
    sub_text: str = Field(default="", description="하단 부제 텍스트 (선택)")
    sub_font_size: int = Field(default=36, description="부제 글자 크기")


class _GenerateAIImageInput(BaseModel):
    prompt: str = Field(description="생성할 이미지 설명 (영어 또는 한국어). 구체적일수록 좋음")
    output_path: str = Field(default="ai_image.jpg", description="저장할 파일 경로 (기본: ai_image.jpg)")
    size: str = Field(default="1024x1024", description="이미지 크기 — 1024x1024(정사각형), 1792x1024(가로), 1024x1792(세로)")
    quality: str = Field(default="standard", description="품질 — standard(빠름), hd(고화질, 2배 비용)")
    style: str = Field(default="vivid", description="스타일 — vivid(선명/생생), natural(자연스러운/사실적)")
    overlay_text: str = Field(default="", description="생성된 이미지 위에 추가로 넣을 글귀 (선택)")
    overlay_font_size: int = Field(default=52, description="글귀 글자 크기 (기본: 52)")


class _InstagramPostInput(BaseModel):
    image_path: str = Field(description="업로드할 이미지 파일 경로 (JPG, PNG)")
    caption: str = Field(description="게시물 본문 캡션")
    hashtags: str = Field(default="", description="해시태그 (예: #일상 #감성)")


class _InstagramCarouselInput(BaseModel):
    image_paths: list[str] = Field(description="업로드할 이미지 경로 목록 (최대 10개). 예: ['slide1.jpg', 'slide2.jpg']")
    caption: str = Field(description="게시물 본문 캡션")
    hashtags: str = Field(default="", description="해시태그 (예: #일상 #감성)")


class _CreateCarouselInput(BaseModel):
    title: str = Field(description="시리즈 제목 (커버 슬라이드에 표시). 예: '소개팅에서 하면 웃는 순간 5'")
    slides: list[str] = Field(description="각 슬라이드의 텍스트 목록. 예: ['1. 메뉴판 보면서 가격 계산', '2. 화장실 가는 척 ...']")
    style: str = Field(default="dark", description="슬라이드 스타일 — dark(어두운), light(밝은), gradient(그라데이션), minimal(미니멀)")
    output_dir: str = Field(default="carousel", description="슬라이드 저장 폴더명 (기본: carousel)")
    use_ai_background: bool = Field(default=False, description="DALL-E 3 AI 배경 이미지 사용 여부 (True시 시간/비용 증가)")
    ai_theme: str = Field(default="", description="AI 배경 테마 (use_ai_background=True일 때). 예: '카페 분위기', '밤하늘'")


class _InstagramHashtagInput(BaseModel):
    topic: str = Field(description="게시물 주제 (예: 일상, 카페, 여행, 음식, 감성)")
    count: int = Field(default=20, description="추천 해시태그 수")


class _EmptyInput(BaseModel):
    pass


def _ig_get_font(size: int, bold: bool = False):
    from PIL import ImageFont
    path = "C:/Windows/Fonts/malgunbd.ttf" if bold else "C:/Windows/Fonts/malgun.ttf"
    try:
        return ImageFont.truetype(path, size)
    except Exception:
        return ImageFont.load_default()


def _ig_wrap_text(text: str, font, max_width: int, draw) -> list[str]:
    words = text.split()
    lines, cur = [], ""
    for word in words:
        test = (cur + " " + word).strip()
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width:
            cur = test
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines if lines else [text]


def _add_text_to_image(image_path: str, text: str, output_path: str = "",
                       font_size: int = 48, text_color: str = "white",
                       position: str = "center", shadow: bool = True,
                       overlay_opacity: int = 40) -> str:
    from PIL import Image, ImageDraw
    if not Path(image_path).exists():
        return f"오류: 이미지 파일을 찾을 수 없습니다 — {image_path}"
    if not output_path:
        p = Path(image_path)
        output_path = str(p.parent / f"{p.stem}_text{p.suffix}")
    img = Image.open(image_path).convert("RGBA")
    W, H = img.size
    if overlay_opacity > 0:
        overlay = Image.new("RGBA", (W, H), (0, 0, 0, int(255 * overlay_opacity / 100)))
        img = Image.alpha_composite(img, overlay)
    draw = ImageDraw.Draw(img)
    font = _ig_get_font(font_size, bold=True)
    padding = int(W * 0.08)
    color_map = {"white": (255,255,255,255), "black": (0,0,0,255),
                 "yellow": (255,230,80,255), "red": (255,80,80,255)}
    if text_color.startswith("#"):
        r,g,b = int(text_color[1:3],16), int(text_color[3:5],16), int(text_color[5:7],16)
        fill = (r, g, b, 255)
    else:
        fill = color_map.get(text_color.lower(), (255,255,255,255))
    raw_lines = text.replace("\\n", "\n").split("\n")
    lines = []
    for raw in raw_lines:
        lines.extend(_ig_wrap_text(raw, font, W - padding * 2, draw))
    line_h = font_size + int(font_size * 0.3)
    total_h = line_h * len(lines)
    start_y = int(H*0.12) if position=="top" else (H - total_h - int(H*0.12) if position=="bottom" else (H-total_h)//2)
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = (W - (bbox[2]-bbox[0])) // 2
        y = start_y + i * line_h
        if shadow:
            off = max(2, font_size//20)
            draw.text((x+off, y+off), line, font=font, fill=(0,0,0,160))
        draw.text((x, y), line, font=font, fill=fill)
    img.convert("RGB").save(output_path, quality=95)
    return f"이미지 텍스트 합성 완료\n저장 경로: {output_path}\n크기: {W}×{H}px"


def _create_text_image(text: str, output_path: str = "instagram_post.jpg",
                       width: int = 1080, height: int = 1080,
                       bg_color: str = "#1A1A2E", text_color: str = "white",
                       font_size: int = 60, sub_text: str = "", sub_font_size: int = 36) -> str:
    from PIL import Image, ImageDraw
    color_map = {"black":"#000000","white":"#FFFFFF","navy":"#1A1A2E",
                 "purple":"#2D1B69","gray":"#2D3748","dark":"#111111"}
    bg_hex = color_map.get(bg_color.lower(), bg_color)
    def hex_to_rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2],16) for i in (0,2,4))
    img = Image.new("RGB", (width, height), hex_to_rgb(bg_hex))
    draw = ImageDraw.Draw(img)
    accent = (100,80,200)
    draw.rectangle([width//2-40, height//2-int(height*0.35), width//2+40, height//2-int(height*0.35)+3], fill=accent)
    font = _ig_get_font(font_size, bold=True)
    padding = int(width * 0.1)
    if text_color.startswith("#"):
        fill = hex_to_rgb(text_color)
    else:
        fill = {"white":(255,255,255),"black":(0,0,0),"yellow":(255,230,80),"gold":(212,175,55)}.get(text_color.lower(),(255,255,255))
    raw_lines = text.replace("\\n", "\n").split("\n")
    lines = []
    for raw in raw_lines:
        lines.extend(_ig_wrap_text(raw, font, width - padding*2, draw))
    line_h = font_size + int(font_size*0.4)
    total_h = line_h * len(lines)
    start_y = (height - total_h) // 2
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0,0), line, font=font)
        x = (width - (bbox[2]-bbox[0])) // 2
        draw.text((x, start_y + i*line_h), line, font=font, fill=fill)
    if sub_text:
        sfont = _ig_get_font(sub_font_size, bold=False)
        sub_fill = (180,180,180)
        sbbox = draw.textbbox((0,0), sub_text, font=sfont)
        draw.text(((width-(sbbox[2]-sbbox[0]))//2, start_y+total_h+int(font_size*0.8)), sub_text, font=sfont, fill=sub_fill)
    img.save(output_path, quality=95)
    return f"이미지 생성 완료\n저장 경로: {output_path}\n크기: {width}×{height}px"


def _generate_ai_image(prompt: str, output_path: str = "ai_image.jpg",
                       size: str = "1024x1024", quality: str = "standard",
                       style: str = "vivid", overlay_text: str = "",
                       overlay_font_size: int = 52) -> str:
    """DALL-E 3로 AI 이미지를 생성하고, 선택적으로 글귀를 합성합니다."""
    import httpx
    from openai import OpenAI

    client_ai = OpenAI(api_key=OPENAI_API_KEY)
    try:
        response = client_ai.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size=size,
            quality=quality,
            style=style,
            n=1,
        )
        image_url = response.data[0].url
        revised_prompt = response.data[0].revised_prompt or prompt

        # URL에서 이미지 다운로드
        img_bytes = httpx.get(image_url, timeout=30).content
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(img_bytes)

        result = (f"AI 이미지 생성 완료\n"
                  f"저장 경로: {output_path}\n"
                  f"크기: {size}\n"
                  f"품질: {quality} / 스타일: {style}\n"
                  f"프롬프트 해석: {revised_prompt[:100]}...")

        # 글귀 합성 (선택)
        if overlay_text:
            merged = _add_text_to_image(
                image_path=output_path,
                text=overlay_text,
                output_path=output_path,
                font_size=overlay_font_size,
                text_color="white",
                position="center",
                shadow=True,
                overlay_opacity=35,
            )
            result += f"\n\n글귀 합성: {merged}"

        return result

    except Exception as e:
        return f"이미지 생성 실패: {e}"


def _create_carousel_images(title: str, slides: list[str], style: str = "dark",
                            output_dir: str = "carousel",
                            use_ai_background: bool = False,
                            ai_theme: str = "") -> str:
    """PPT 슬라이드처럼 여러 장의 카드 이미지를 생성합니다."""
    from PIL import Image, ImageDraw
    import httpx
    from openai import OpenAI

    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    STYLES = {
        "dark":     {"bg": (18, 18, 30),    "accent": (120, 80, 220), "text": (255,255,255), "sub": (180,180,200)},
        "light":    {"bg": (248, 248, 248),  "accent": (80, 80, 220),  "text": (20, 20, 40),  "sub": (100,100,120)},
        "gradient": {"bg": (15, 15, 35),     "accent": (200, 80, 180), "text": (255,255,255), "sub": (200,180,220)},
        "minimal":  {"bg": (255, 255, 255),  "accent": (40, 40, 40),   "text": (30, 30, 30),  "sub": (120,120,120)},
    }
    s = STYLES.get(style, STYLES["dark"])
    W, H = 1080, 1080
    saved = []

    def draw_slide(idx: int, headline: str, body: str = "",
                   is_cover: bool = False, bg_img_bytes: bytes = None) -> str:
        if bg_img_bytes:
            import io
            base = Image.open(io.BytesIO(bg_img_bytes)).convert("RGBA").resize((W, H))
            overlay = Image.new("RGBA", (W, H), (0, 0, 0, 170))
            img = Image.alpha_composite(base, overlay).convert("RGB")
        else:
            img = Image.new("RGB", (W, H), s["bg"])

        draw = ImageDraw.Draw(img)

        # 번호 뱃지
        if not is_cover:
            badge_r = 38
            bx, by = 60, 60
            draw.ellipse([bx, by, bx+badge_r*2, by+badge_r*2], fill=s["accent"])
            nfont = _ig_get_font(28, bold=True)
            nb = draw.textbbox((0,0), str(idx), font=nfont)
            draw.text((bx + badge_r - (nb[2]-nb[0])//2,
                       by + badge_r - (nb[3]-nb[1])//2),
                      str(idx), font=nfont, fill=(255,255,255))

        # 액센트 라인
        if is_cover:
            lw = 80
            draw.rectangle([W//2 - lw//2, H//2 - 160, W//2 + lw//2, H//2 - 155], fill=s["accent"])
        else:
            draw.rectangle([60, H - 70, W - 60, H - 65], fill=(*s["accent"], 120))

        # 헤드라인 텍스트
        if is_cover:
            fsize = 64
            font = _ig_get_font(fsize, bold=True)
            raw = headline.replace("\\n", "\n").split("\n")
            lines = []
            for r in raw:
                lines.extend(_ig_wrap_text(r, font, W - 160, draw))
            lh = fsize + int(fsize * 0.45)
            total_h = lh * len(lines)
            sy = (H - total_h) // 2 - 20
            for i, line in enumerate(lines):
                bb = draw.textbbox((0,0), line, font=font)
                x = (W - (bb[2]-bb[0])) // 2
                draw.text((x+2, sy+i*lh+2), line, font=font, fill=(0,0,0,100))
                draw.text((x, sy+i*lh), line, font=font, fill=s["text"])
            # 부제 "Swipe →"
            sf = _ig_get_font(30)
            sb = draw.textbbox((0,0), "Swipe →", font=sf)
            draw.text(((W-(sb[2]-sb[0]))//2, sy+total_h+50), "Swipe →", font=sf, fill=s["sub"])
        else:
            fsize = 54
            font = _ig_get_font(fsize, bold=True)
            raw = headline.replace("\\n", "\n").split("\n")
            lines = []
            for r in raw:
                lines.extend(_ig_wrap_text(r, font, W - 160, draw))
            lh = fsize + int(fsize * 0.4)
            total_h = lh * len(lines)
            sy = (H - total_h) // 2
            for i, line in enumerate(lines):
                bb = draw.textbbox((0,0), line, font=font)
                x = (W - (bb[2]-bb[0])) // 2
                draw.text((x+2, sy+i*lh+2), line, font=font, fill=(0,0,0,120))
                draw.text((x, sy+i*lh), line, font=font, fill=s["text"])

            if body:
                bfont = _ig_get_font(32)
                braw = body.replace("\\n", "\n").split("\n")
                blines = []
                for r in braw:
                    blines.extend(_ig_wrap_text(r, bfont, W - 200, draw))
                bsy = sy + total_h + 40
                for i, line in enumerate(blines):
                    bb = draw.textbbox((0,0), line, font=bfont)
                    x = (W - (bb[2]-bb[0])) // 2
                    draw.text((x, bsy + i*50), line, font=bfont, fill=s["sub"])

        path = str(out / f"slide_{idx:02d}.jpg")
        img.save(path, quality=95)
        return path

    # AI 배경 공통 생성 (use_ai_background=True)
    ai_bg_bytes = None
    if use_ai_background and ai_theme:
        try:
            cl_ai = OpenAI(api_key=OPENAI_API_KEY)
            resp = cl_ai.images.generate(
                model="dall-e-3",
                prompt=f"{ai_theme}, cinematic, moody, Instagram aesthetic, 1:1 square",
                size="1024x1024", quality="standard", style="vivid", n=1,
            )
            ai_bg_bytes = httpx.get(resp.data[0].url, timeout=30).content
        except Exception:
            ai_bg_bytes = None

    # 커버 슬라이드 (index 0)
    path = draw_slide(0, title, is_cover=True, bg_img_bytes=ai_bg_bytes)
    saved.append(path)

    # 콘텐츠 슬라이드 (index 1~)
    for i, text in enumerate(slides, start=1):
        path = draw_slide(i, text, is_cover=False, bg_img_bytes=ai_bg_bytes)
        saved.append(path)

    return (f"캐러셀 슬라이드 {len(saved)}장 생성 완료\n"
            f"저장 폴더: {output_dir}/\n"
            f"파일 목록: {saved}")


def _instagram_post_carousel(image_paths: list[str], caption: str, hashtags: str = "") -> str:
    """여러 이미지를 Instagram 캐러셀(슬라이드) 게시물로 포스팅합니다."""
    from instagram_mcp_server import _check_instagram_config, _get_instagram_client

    err = _check_instagram_config()
    if err:
        return err

    missing = [p for p in image_paths if not Path(p).exists()]
    if missing:
        return f"오류: 다음 파일을 찾을 수 없습니다 — {missing}"
    if len(image_paths) < 2:
        return "캐러셀은 이미지 2장 이상 필요합니다. 단일 이미지는 instagram_post를 사용하세요."
    if len(image_paths) > 10:
        return "Instagram 캐러셀은 최대 10장까지 가능합니다."

    full_caption = f"{caption}\n\n{hashtags}" if hashtags else caption
    try:
        cl = _get_instagram_client()
        media = cl.album_upload(paths=image_paths, caption=full_caption)
        return (f"Instagram 캐러셀 포스팅 완료!\n"
                f"게시물 ID: {media.pk}\n"
                f"슬라이드 수: {len(image_paths)}장\n"
                f"캡션: {caption[:60]}{'...' if len(caption)>60 else ''}")
    except Exception as e:
        msg = str(e)
        if "challenge_required" in msg.lower():
            return "보안 인증 필요 — 인스타그램 앱에서 로그인 알림을 승인해 주세요."
        return f"포스팅 실패: {msg}"


def _instagram_post(image_path: str, caption: str, hashtags: str = "") -> str:
    from instagram_mcp_server import _check_instagram_config, _get_instagram_client
    err = _check_instagram_config()
    if err:
        return err
    if not Path(image_path).exists():
        return f"오류: 이미지 파일을 찾을 수 없습니다 — {image_path}"
    full_caption = f"{caption}\n\n{hashtags}" if hashtags else caption
    try:
        cl = _get_instagram_client()
        media = cl.photo_upload(path=image_path, caption=full_caption)
        return (f"Instagram 포스팅 완료!\n게시물 ID: {media.pk}\n"
                f"캡션: {caption[:50]}{'...' if len(caption)>50 else ''}")
    except Exception as e:
        msg = str(e)
        if "challenge_required" in msg.lower():
            return "보안 인증 필요 — 인스타그램 앱에서 로그인 알림을 승인해 주세요."
        return f"포스팅 실패: {msg}"


def _instagram_check_login() -> str:
    from instagram_mcp_server import _check_instagram_config, _get_instagram_client
    err = _check_instagram_config()
    if err:
        return err
    try:
        cl = _get_instagram_client()
        user = cl.account_info()
        return (f"Instagram 로그인 성공\n사용자명: @{user.username}\n"
                f"이름: {user.full_name or '(미설정)'}\n이메일: {user.email or '(미설정)'}")
    except Exception as e:
        return f"로그인 실패: {e}"


def _instagram_generate_hashtags(topic: str, count: int = 20) -> str:
    tag_db = {
        "일상": ["#일상","#daily","#오늘","#vlog","#일상스타그램","#데일리","#소소한일상","#일상공유","#맞팔","#좋아요"],
        "카페": ["#카페","#cafe","#커피","#coffee","#카페스타그램","#카페투어","#카페인","#핸드드립","#라떼아트","#디저트"],
        "여행": ["#여행","#travel","#여행스타그램","#旅行","#trip","#여행사진","#국내여행","#해외여행","#풍경","#감성여행"],
        "음식": ["#맛스타그램","#먹스타그램","#food","#맛집","#foodporn","#맛집탐방","#홈쿡","#요리","#delicious","#yummy"],
        "감성": ["#감성","#감성사진","#감성스타그램","#mood","#aesthetic","#감성충전","#필름사진","#사진스타그램","#포토그래피","#인생사진"],
        "자연": ["#자연","#nature","#풍경","#landscape","#하늘","#숲","#바다","#산","#힐링","#자연스타그램"],
    }
    result = []
    for key, tags in tag_db.items():
        if key in topic or topic in key:
            result.extend(tags)
    if not result:
        result = ["#일상","#daily","#감성","#사진","#photo","#스타그램","#인스타","#좋아요","#맞팔","#follow"]
    result = list(dict.fromkeys(result))[:count]
    return f"추천 해시태그 ({len(result)}개):\n" + " ".join(result)


INSTAGRAM_TOOLS = [
    StructuredTool.from_function(
        func=_generate_ai_image, name="generate_ai_image",
        description=(
            "DALL-E 3 AI로 실제 이미지를 생성합니다. "
            "원하는 장면, 분위기, 스타일을 프롬프트로 설명하면 고품질 이미지를 만듭니다. "
            "overlay_text로 생성된 이미지 위에 글귀도 합성할 수 있습니다. "
            "Instagram 포스팅용 이미지 제작에 적합합니다."
        ),
        args_schema=_GenerateAIImageInput,
    ),
    StructuredTool.from_function(
        func=_add_text_to_image, name="add_text_to_image",
        description="기존 사진 위에 글귀/텍스트를 합성합니다. 그림자·오버레이 자동 처리.",
        args_schema=_AddTextToImageInput,
    ),
    StructuredTool.from_function(
        func=_create_text_image, name="create_text_image",
        description="단색 배경 위에 글귀만으로 감성 이미지를 생성합니다. AI 이미지가 필요 없을 때 사용.",
        args_schema=_CreateTextImageInput,
    ),
    StructuredTool.from_function(
        func=_create_carousel_images, name="create_carousel_images",
        description=(
            "PPT 슬라이드처럼 여러 장의 카드 이미지를 생성합니다. "
            "title(커버), slides(각 슬라이드 텍스트 목록)을 입력하면 "
            "번호 뱃지와 감성 레이아웃이 자동 적용된 슬라이드 이미지를 만듭니다. "
            "style: dark/light/gradient/minimal 선택 가능. "
            "use_ai_background=True로 DALL-E 3 배경도 사용 가능."
        ),
        args_schema=_CreateCarouselInput,
    ),
    StructuredTool.from_function(
        func=_instagram_post_carousel, name="instagram_post_carousel",
        description=(
            "여러 이미지(최대 10장)를 Instagram 캐러셀(슬라이드) 게시물로 포스팅합니다. "
            "image_paths에 파일 경로 목록을 전달하세요. "
            "create_carousel_images로 슬라이드를 생성한 뒤 이 도구로 업로드합니다."
        ),
        args_schema=_InstagramCarouselInput,
    ),
    StructuredTool.from_function(
        func=_instagram_post, name="instagram_post",
        description="이미지 + 캡션 + 해시태그로 Instagram에 포스팅합니다.",
        args_schema=_InstagramPostInput,
    ),
    StructuredTool.from_function(
        func=_instagram_check_login, name="instagram_check_login",
        description="Instagram 로그인 상태와 계정 정보를 확인합니다.",
        args_schema=_EmptyInput,
    ),
    StructuredTool.from_function(
        func=_instagram_generate_hashtags, name="instagram_generate_hashtags",
        description="주제에 맞는 인기 해시태그 조합을 추천합니다.",
        args_schema=_InstagramHashtagInput,
    ),
]


# ─────────────────────────────────────────
# 비동기 실행 유틸
# ─────────────────────────────────────────

def run_async(coro, timeout: int = 180) -> str:
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        return ex.submit(asyncio.run, coro).result(timeout=timeout)


async def _invoke(history: list[dict]) -> str:
    # 모든 도구를 직접(in-process) 로드 — MCP 서브프로세스 없음
    all_tools = (EXCEL_TOOLS + UTILS_TOOLS + EMAIL_TOOLS
                 + RAG_TOOLS + SQL_TOOLS + DOCUMENT_TOOLS + INSTAGRAM_TOOLS)

    selected_provider = st.session_state.get("model_provider", MODEL_PROVIDER)
    selected_local    = st.session_state.get("local_model_choice", LOCAL_MODEL)
    if selected_provider == "local":
        llm = ChatOpenAI(
            model=selected_local,
            base_url=OLLAMA_BASE_URL,
            api_key="ollama",
            temperature=0,
        )
    else:
        llm = build_llm("openai")
    agent = create_react_agent(llm, all_tools, prompt=SYSTEM_PROMPT)

    messages = []
    for m in history:
        if m["role"] == "user":
            messages.append(HumanMessage(content=m["content"]))
        elif m["role"] == "assistant":
            messages.append(AIMessage(content=m["content"]))

    result = await agent.ainvoke({"messages": messages})
    for msg in reversed(result["messages"]):
        if isinstance(msg, AIMessage) and msg.content:
            return msg.content
    return "응답을 생성하지 못했습니다."


def get_response(history: list[dict]) -> str:
    return run_async(_invoke(history))


def extract_image_paths(text: str) -> list[str]:
    found = re.findall(r'[A-Za-z0-9가-힣_\-. /\\:()]+\.png', text)
    return [p.strip() for p in found if Path(p.strip()).exists()]


# ─────────────────────────────────────────
# 페이지 설정
# ─────────────────────────────────────────
st.set_page_config(
    page_title="AI Agent",
    page_icon=None,
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────
# CSS — 프로페셔널 UI
# ─────────────────────────────────────────
st.markdown("""
<style>
html, body, [class*="css"] {
    font-family: "Pretendard", "Noto Sans KR", -apple-system, sans-serif;
}
.stApp { background: #F7F8FA; }
#MainMenu, footer, header { visibility: hidden; }

section[data-testid="stSidebar"] > div:first-child {
    background: #16161A;
    padding-top: 1.5rem;
}
section[data-testid="stSidebar"] * { color: #E2E8F0 !important; }
section[data-testid="stSidebar"] .stDivider { border-color: #2D2D35 !important; }
section[data-testid="stSidebar"] code {
    background: #2D2D35 !important;
    color: #A0AEC0 !important;
    border-radius: 4px;
    font-size: 0.78rem !important;
}
section[data-testid="stSidebar"] button {
    background: #2D2D35 !important;
    color: #E2E8F0 !important;
    border: 1px solid #3D3D4A !important;
    border-radius: 6px !important;
    font-size: 0.82rem !important;
    transition: background 0.2s;
}
section[data-testid="stSidebar"] button:hover { background: #3D3D4A !important; }

.main .block-container { max-width: 820px; padding: 2rem 1.5rem 1rem; }

[data-testid="stChatMessage"] {
    background: white;
    border: 1px solid #E8ECF0;
    border-radius: 10px;
    margin-bottom: 0.6rem;
    padding: 0.9rem 1.1rem;
    box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}
[data-testid="stChatMessage"][data-role="user"] {
    background: #1A1A2E;
    border-color: #1A1A2E;
}
[data-testid="stChatMessage"][data-role="user"] p,
[data-testid="stChatMessage"][data-role="user"] span { color: #F0F0F5 !important; }

[data-testid="stChatInput"] {
    border: 1.5px solid #D1D5DB !important;
    border-radius: 10px !important;
    background: white !important;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
}
[data-testid="stChatInput"]:focus-within {
    border-color: #553C9A !important;
    box-shadow: 0 0 0 3px rgba(85,60,154,0.1) !important;
}
.stSpinner > div { border-top-color: #553C9A !important; }
code {
    background: #EEF2F8 !important;
    color: #2D3748 !important;
    border-radius: 4px;
    font-size: 0.83rem !important;
}
h1 { color: #1A1A2E; font-weight: 700; letter-spacing: -0.5px; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────
# 사이드바
# ─────────────────────────────────────────
with st.sidebar:
    st.markdown("### AI Agent")

    # ── 모델 선택 ──
    OLLAMA_MODELS = ["deepseek-r1:8b", "qwen2.5:7b-instruct", "llama3.2:1b"]

    if "model_provider" not in st.session_state:
        st.session_state["model_provider"] = MODEL_PROVIDER
    if "local_model_choice" not in st.session_state:
        st.session_state["local_model_choice"] = LOCAL_MODEL

    provider_choice = st.radio(
        "모델 프로바이더",
        ["OpenAI", "Local (Ollama)"],
        index=0 if st.session_state["model_provider"] == "openai" else 1,
        key="_provider_radio",
        horizontal=True,
    )
    new_provider = "openai" if provider_choice == "OpenAI" else "local"

    if new_provider == "local":
        local_choice = st.selectbox(
            "로컬 모델 선택",
            OLLAMA_MODELS,
            index=OLLAMA_MODELS.index(st.session_state["local_model_choice"])
                  if st.session_state["local_model_choice"] in OLLAMA_MODELS else 0,
            key="_local_model_select",
        )
        if local_choice != st.session_state["local_model_choice"]:
            st.session_state["local_model_choice"] = local_choice

    if new_provider != st.session_state["model_provider"]:
        st.session_state["model_provider"] = new_provider
        st.rerun()

    active_model = (OPENAI_MODEL if st.session_state["model_provider"] == "openai"
                    else st.session_state["local_model_choice"])
    st.markdown(
        f"<span style='font-size:0.78rem;color:#888;'>현재: <code>{active_model}</code> · LangGraph + MCP</span>",
        unsafe_allow_html=True,
    )
    if st.session_state["model_provider"] == "local":
        st.info("Ollama 로컬 실행 중 — 인터넷 불필요, 무료", icon="💻")

    st.divider()

    st.markdown("**Excel 도구**")
    st.markdown("""
- 파일 생성 / 읽기
- 셀 수정 · 수식 적용
- 시트 관리 (추가 · 이름변경)
- 데이터 일괄 입력 · 행 삭제
    """)

    st.markdown("**유틸리티 도구**")
    st.markdown("""
- 웹 검색 (DuckDuckGo)
- 현재 날짜 / 시각 조회
- 디렉토리 파일 목록
- 텍스트 · CSV · JSON 읽기
- 차트 생성 (막대 · 선 · 원형)
    """)

    st.markdown("**해양 도메인 RAG**")
    st.markdown("""
- OSP Interface Specification
- FMI 2.0 표준
- 동적 위치 유지 (DP)
- 크레인 선박 / 해양 건설
- 해양 전력 관리
    """)

    st.markdown("**이메일 도구**")
    gmail = os.getenv("GMAIL_EMAIL", "미설정")
    app_pw_set = "설정됨" if os.getenv("GMAIL_APP_PASSWORD") else "미설정"
    st.markdown(f"""
- 계정: `{gmail}`
- 앱 비밀번호: `{app_pw_set}`
- 텍스트 이메일 발송
- 파일 첨부 이메일 (Excel · PDF · PNG)
- 참조(CC) 지원
    """)

    st.markdown("**Instagram 자동화**")
    insta_user = os.getenv("INSTAGRAM_USERNAME", "")
    insta_status = f"`@{insta_user}`" if insta_user else "`미설정`"
    st.markdown(f"""
- 계정: {insta_status}
- DALL-E 3 AI 이미지 생성
- 사진 위 글귀 합성
- 캡션 + 해시태그 추천
- 단일 이미지 포스팅
- **캐러셀(슬라이드) 포스팅** (최대 10장)
    """)

    st.markdown("**SQL 데이터베이스**")
    st.markdown("""
- SELECT 쿼리 · 집계 분석
- 테이블 생성 · 데이터 입력
- 스키마 조회
- Excel → DB 변환
    """)

    st.markdown("**문서 처리 (PDF / Word)**")
    st.markdown("""
- PDF 텍스트 추출 (페이지 범위 지정)
- Word(.docx) 텍스트 · 표 추출
- 문서 메타데이터 조회
    """)

    st.divider()
    st.markdown("**예시**")
    for ex in [
        "1~12월 매출 엑셀 만들어줘",
        "매출 데이터로 막대 차트 만들어줘",
        "DP 시스템이란 무엇인가?",
        "test@gmail.com에 '안녕' 메일 보내줘",
        "매출 테이블 만들고 데이터 넣어줘",
        "report.pdf 내용 요약해줘",
    ]:
        st.code(ex, language=None)

    st.divider()
    if st.button("대화 초기화", use_container_width=True):
        st.session_state.messages = []
        st.rerun()


# ─────────────────────────────────────────
# 메인 헤더
# ─────────────────────────────────────────
st.markdown("## AI Agent")
st.markdown(
    "<p style='color:#718096;font-size:0.9rem;margin-top:-0.8rem;'>"
    "Excel · SQL · 문서(PDF/Word) · 웹 검색 · 이메일 · 해양 도메인 RAG"
    "</p>",
    unsafe_allow_html=True
)
st.divider()


# ─────────────────────────────────────────
# 채팅 히스토리 초기화
# ─────────────────────────────────────────
if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": (
                "안녕하세요. 무엇을 도와드릴까요?\n\n"
                "다음 작업을 처리할 수 있습니다.\n\n"
                "- **Excel** 파일 작성·수정·수식 적용\n"
                "- **SQL** 데이터베이스 조회·분석·Excel → DB 변환\n"
                "- **PDF / Word** 문서 텍스트 추출·요약\n"
                "- **웹 검색** 및 차트·파일 관리\n"
                "- **이메일** 발송 (첨부 파일 포함)\n"
                "- **해양 기술** (DP, OSP, FMI 등) 지식베이스 검색"
            ),
        }
    ]

# 이전 메시지 렌더링
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        for img_path in extract_image_paths(msg["content"]):
            st.image(img_path, use_container_width=True)


# ─────────────────────────────────────────
# 사용자 입력 처리
# ─────────────────────────────────────────
if prompt := st.chat_input("요청 사항을 입력하세요..."):
    with st.chat_message("user"):
        st.markdown(prompt)
    st.session_state.messages.append({"role": "user", "content": prompt})

    with st.chat_message("assistant"):
        with st.spinner("처리 중..."):
            try:
                response = get_response(st.session_state.messages)
            except concurrent.futures.TimeoutError:
                response = "응답 시간이 초과되었습니다. 다시 시도해 주세요."
            except Exception as e:
                response = f"오류가 발생했습니다: {str(e)}"

        st.markdown(response)
        for img_path in extract_image_paths(response):
            st.image(img_path, use_container_width=True)

    st.session_state.messages.append({"role": "assistant", "content": response})
